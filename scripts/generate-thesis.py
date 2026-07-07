#!/usr/bin/env python3
"""Generate BIT-SE Bachelor's Thesis document for BIU Lost & Found project."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


AUTHOR = "NHANH KIMSON"
YEAR = "2026"
SUPERVISOR = "MR. SOTHEA THY"
TITLE = (
    "BUILD A DIGITAL LOST & FOUND SYSTEM FOR BELTEI "
    "INTERNATIONAL UNIVERSITY"
)
SHORT_TITLE = "Build a Digital Lost & Found System"
SYSTEM_ACRONYM = "DLFS"
ORG = "BELTEI International University"
SAMPLE_SIZE = 230
GOOGLE_FORM_URL = (
    "https://docs.google.com/forms/d/e/1FAIpQLSdUNzeq_a1WyzGIk1rZ3T352yGHpd5ZQ3jbm-LdWPrgUwk-4Q/viewform"
)
OUTPUT = Path(__file__).resolve().parents[1] / "thesis" / "Nhanh-Kimson-Thesis-2026.docx"
OUTPUT_APPLICATIONS = Path("/Applications/Nhanh Kimson.docx")
CHAPTER3_TXT = Path(__file__).resolve().parents[1] / "thesis" / "chapter-3.txt"
CHAPTER4_TXT = Path(__file__).resolve().parents[1] / "thesis" / "chapter-4.txt"
FIGURE_1 = Path(__file__).resolve().parents[1] / "thesis" / "figures" / "figure-1-conceptual-framework.png"

import sys

_THESIS_DIR = Path(__file__).resolve().parents[1] / "thesis"
if str(_THESIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THESIS_DIR))
from chapter_4_sections import (  # noqa: E402
    CHAPTER4_INTRO,
    EXAMINER_SUMMARY,
    HYPOTHESES_CLOSING,
    HYPOTHESES_DETAILED,
    HYPOTHESES_INTRO,
    HYPOTHESIS_TABLE,
    SECTION_4_1_INTRO,
    SECTION_4_2_INTRO,
    SECTION_4_3_INTRO,
    SOLUTION_BLOCKS,
    STRENGTHS_BLOCKS,
    WEAKNESSES_BLOCKS,
)


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_centered(doc: Document, text: str, *, bold: bool = False, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = None


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_figure(doc: Document, path: Path, caption: str, *, width: float = 5.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_centered(doc, caption)


def ensure_thesis_figures() -> None:
    import subprocess
    import sys

    script = Path(__file__).parent / "generate-framework-diagram.py"
    subprocess.run([sys.executable, str(script)], check=True)


def _table_to_text(table) -> str:
    lines: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def export_chapter3_txt(doc: Document) -> None:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    header = (
        "================================================================================\n"
        "UPDATED — synced from scripts/generate-thesis.py\n"
        f"System: {SYSTEM_ACRONYM} | Sample: {SAMPLE_SIZE} respondents | Author: {AUTHOR}\n"
        "================================================================================\n"
    )
    lines: list[str] = [header.rstrip(), ""]
    in_ch3 = False

    for child in doc.element.body:
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text == "CHAPTER 3: RESEARCH METHODOLOGY":
                in_ch3 = True
            if in_ch3 and text in ("CHAPTER 4: DATA ANALYSIS", "CHAPTER 4 DATA ANALYSIS"):
                break
            if in_ch3 and text:
                lines.append(text)
                lines.append("")
        elif isinstance(child, CT_Tbl) and in_ch3:
            lines.append(_table_to_text(Table(child, doc)))
            lines.append("")

    CHAPTER3_TXT.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    print(f"Chapter 3 exported: {CHAPTER3_TXT}")


def export_chapter4_txt(doc: Document) -> None:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    header = (
        "================================================================================\n"
        "UPDATED — synced from scripts/generate-thesis.py\n"
        f"System: {SYSTEM_ACRONYM} | Sample: {SAMPLE_SIZE} respondents | Author: {AUTHOR}\n"
        "================================================================================\n"
    )
    lines: list[str] = [header.rstrip(), ""]
    in_ch4 = False

    for child in doc.element.body:
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text == "CHAPTER 4 DATA ANALYSIS" or text == "CHAPTER 4: DATA ANALYSIS":
                in_ch4 = True
            if in_ch4 and (
                text.startswith("CHAPTER 5")
                or text == "CHAPTER 5: RESEARCH FINDINGS AND DISCUSSIONS"
            ):
                break
            if in_ch4 and text:
                lines.append(text)
                lines.append("")
        elif isinstance(child, CT_Tbl) and in_ch4:
            lines.append(_table_to_text(Table(child, doc)))
            lines.append("")

    CHAPTER4_TXT.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    print(f"Chapter 4 exported: {CHAPTER4_TXT}")


def build_front_matter(doc: Document) -> None:
    for _ in range(3):
        add_centered(doc, TITLE, bold=True, size=14)
        doc.add_paragraph()
    add_centered(doc, AUTHOR, bold=True)
    add_centered(doc, YEAR, bold=True)
    add_page_break(doc)

    add_centered(doc, TITLE, bold=True, size=14)
    doc.add_paragraph()
    add_centered(
        doc,
        "THE BACHELOR'S THESIS IN PARTIAL FULFILLMENT\n"
        "OF THE REQUIREMENT FOR THE DEGREE OF BACHELOR\n"
        "OF INFORMATION TECHNOLOGY IN SOFTWARE ENGINEERING",
        size=12,
    )
    add_page_break(doc)

    add_centered(doc, TITLE, bold=True, size=14)
    doc.add_paragraph()
    add_centered(doc, SUPERVISOR)
    add_centered(doc, YEAR)
    add_centered(doc, AUTHOR)
    add_page_break(doc)

    add_centered(doc, "COMMITTEE APPROVAL", bold=True)
    doc.add_paragraph()
    add_body(
        doc,
        f'This Bachelor\'s Thesis entitled "{SHORT_TITLE} for {ORG}" was prepared '
        f"and submitted by {AUTHOR} of the Beltei International University in partial "
        "fulfilment of the requirement of a Bachelor of Software Engineering (BIT-SE).",
    )
    add_body(doc, "Approved by the University Evaluation Committee with a grade of PASSED")
    add_body(doc, "The university EVALUATION COMMITTEE OF BACHELOR'S THESIS")
    add_body(doc, "Chair of Committee\t\t:_____________________________")
    add_body(doc, "\t\t\t\tMr. OEM Chanthorn, Dean")
    add_body(doc, "Deputy Chair of Committee\t:_____________________________")
    add_body(doc, "\t\t\t\tMr. KROENG Vannak, Vice Dean")
    add_body(doc, "Member\t\t\t\t:_____________________________")
    add_body(doc, "\t\t\t\tMr. KHATH Yoeun, Master")
    add_body(doc, "Member\t\t\t\t:_____________________________")
    add_body(doc, "\t\t\t\tMr. PRUM Vannarith, Master")
    add_body(doc, "Facilitator\t\t\t:_____________________________")
    add_body(doc, f"\t\t\t\t{SUPERVISOR}, Master")
    add_body(doc, "Date of Bachelor Thesis submission\t:_____________________________")
    add_body(
        doc,
        "Accepted in partial fulfilment of the requirements for the degree of "
        "Bachelor of Information Technology in Software Engineering (BIT-SE)",
    )
    doc.add_paragraph()
    add_centered(doc, "________________________")
    add_centered(doc, "H.E. LY Navuth")
    add_centered(doc, "President")
    add_page_break(doc)

    add_centered(doc, "DECLARATION", bold=True)
    doc.add_paragraph()
    add_body(
        doc,
        f'I do hereby declare that, except otherwise stated the Bachelor\'s Thesis '
        f'"{SHORT_TITLE} for {ORG}" based on my original work and the same has not '
        "been submitted either in part or in full for the award of any other degree "
        "of this or any other University.",
    )
    doc.add_paragraph()
    add_body(doc, "My indebtedness to other writer(s) has/have been acknowledged at relevant places.")
    doc.add_paragraph()
    add_centered(doc, "__________________________                          __________________________")
    add_centered(doc, f"{AUTHOR}                                                                 Date Signed")
    add_page_break(doc)

    add_centered(doc, "ACKNOWLEDGMENT", bold=True)
    doc.add_paragraph()
    add_body(
        doc,
        "I would like to acknowledge people whose contributions helped me to successfully "
        "complete this research report.",
    )
    add_body(
        doc,
        "First, I would like to express my grateful thanks to H.E LY Navuth, the President "
        "of Beltei International University for offering me a great opportunity and scholarship "
        "to study at Beltei International University, in Faculty of Information Technology and "
        "Science, majoring in Software Engineering.",
    )
    add_body(
        doc,
        "Secondly, I would like to express thanks to Mr. OEM CHANTHORN, Dean of Faculty of "
        "Information Technology and Science and Faculty of Digital Technology and Telecommunication "
        "for providing us with the opportunity to join this program of research to gain more "
        "knowledge and improve research skills which is important for academics.",
    )
    add_body(
        doc,
        f"Third, I would like to thank Mr. SOTHEA THY, my supervisor who helped me to supervise "
        "this research report to have a better outcome. Throughout the process of conducting "
        "research, I would like to express most gratefulness for his patience, support, goodwill, "
        "and understanding since it had assisted me excessively for the progress of my research. "
        "He has played a very important role as I went through the steps of this study as he "
        "extended his guidance and perseverance by advising the best choices to make from the "
        "beginning until the end of this research process.",
    )
    add_body(
        doc,
        "Lastly, I deeply extend my appreciation towards the participants of my survey questionnaire "
        "for the essential data collection of this study. Their valuable time, honesty, patience, "
        "and cooperativeness are greatly appreciated as their contribution would yield the best "
        "outcome for this research.",
    )
    add_page_break(doc)

    add_centered(doc, "ABSTRACT", bold=True)
    doc.add_paragraph()
    add_body(
        doc,
        f"Lost and found management is a critical campus service that directly affects student "
        f"satisfaction, operational efficiency, and the safe return of personal belongings. "
        f"For {ORG}, this study investigates the development of a {SHORT_TITLE} ({SYSTEM_ACRONYM}) "
        f"with the goal of improving the speed, transparency, and reliability of reporting lost "
        f"items, publishing found items, matching potential owners, and processing ownership claims. "
        f"Many universities still rely on bulletin boards, informal social media posts, or fragmented "
        f"manual records, which frequently result in delayed reunions, duplicate reports, and poor "
        f"visibility of open cases. This study examines the challenges faced by students and campus "
        f"staff in lost-and-found coordination and proposes a digital solution that integrates "
        f"structured item reporting, image uploads, searchable listings, automated match suggestions, "
        f"claim verification, in-app notifications, and role-based administration.",
    )
    add_body(
        doc,
        "Through a mixed-method approach combining surveys, interviews, and direct observation of "
        "existing campus practices, the study identifies key weaknesses in manual coordination and "
        "evaluates user readiness for a centralized web platform. The proposed system is designed "
        "with a user-friendly interface, secure authentication, validated forms, and a scalable "
        "PostgreSQL database to ensure accuracy, accessibility, and maintainability. Built with "
        "Next.js, Prisma, NextAuth, Cloudinary, and REST APIs documented in OpenAPI, the system "
        "supports both browser users and external clients.",
    )
    add_body(
        doc,
        "The findings indicate that implementing DLFS can significantly reduce the time required "
        "to locate matching listings, improve claim traceability, and increase confidence in campus "
        "lost-and-found processes. The system also supports better decision-making for administrators "
        "through dashboards, claim review tools, and activity analytics. Overall, the research "
        "demonstrates that digital transformation in campus lost-and-found services is both "
        "technically feasible and strongly supported by end users.",
    )
    add_body(
        doc,
        f"Keyword: Lost and Found, Campus Information System, Digital Transformation, "
        f"Web Application, Claim Management, {ORG}, Software Engineering.",
    )
    add_page_break(doc)


def build_contents(doc: Document) -> None:
    add_centered(doc, "CONTENTS", bold=True)
    doc.add_paragraph()
    contents = [
        ("COMMITTEE APPROVAL", "iv"),
        ("DECLARATION", "v"),
        ("ACKNOWLEDGMENT", "vi"),
        ("ABSTRACT", "vii"),
        ("CONTENTS", "viii"),
        ("LIST OF TABLES", "xi"),
        ("LIST OF FIGURES", "xii"),
        ("LIST OF ABBREVIATIONS", "xiii"),
        ("CHAPTER 1: INTRODUCTION", "1"),
        ("1.1 Introduction to Research", "1"),
        ("1.2 Research Problem", "2"),
        ("1.3 Research Aim/Objective", "2"),
        ("1.3.1 Research Question", "3"),
        ("1.3.2 Signification of the study", "3"),
        ("1.4 Scope and limitation", "3"),
        ("1.4.1 Scope and Limitations", "4"),
        ("1.4.2 Limitations of the Study", "4"),
        ("1.5 Layout of the study/Research", "5"),
        ("CHAPTER 2: LITERATURE REVIEW", "7"),
        ("2.1 Definitions and Theories Related to the Topic", "7"),
        ("2.2 Analysis of Scholars' Concepts/Words/Sayings", "8"),
        ("2.3 Experiences or Issues Raised", "9"),
        ("2.4 Research Model or Framework", "10"),
        ("CHAPTER 3: RESEARCH METHODOLOGY", "12"),
        ("3.1 Research Methodology", "12"),
        ("3.1.1 Research Design", "12"),
        ("3.1.2 Research Location", "12"),
        ("3.1.3 Research Population", "14"),
        ("3.1.4 Research Data Scope of Database", "15"),
        ("3.1.5 Research Data Scope of UI", "17"),
        ("3.1.6 Research Data Scope of UX", "21"),
        ("3.2 Data Collection Instrument", "20"),
        ("3.2.1 Data Collection Procedure", "20"),
        ("3.2.2 Statistical Data", "21"),
        ("3.3 Sampling Technique", "31"),
        ("3.4 Validity and Reliability", "32"),
        ("CHAPTER 4: DATA ANALYSIS", "34"),
        ("4.1 Analysis of the Strengths", "34"),
        ("4.1.1 Effectiveness of the Current Lost and Found Practices", "34"),
        ("4.1.2 Communication and Awareness Initiatives", "35"),
        ("4.1.3 Monitoring and Reporting Practices", "36"),
        ("4.2 Analysis of Weaknesses", "37"),
        ("4.2.1 Dependence on Manual and Fragmented Processes", "37"),
        ("4.2.2 Poor Searchability and Incomplete Item Details", "38"),
        ("4.2.3 Delayed Communication and Weak Claim Verification", "39"),
        ("4.3 Solution(s) Dealing with the Weaknesses", "40"),
        ("4.3.1 Implementation of DLFS", "40"),
        ("4.3.2 Integration, Training, and Real-Time Reporting", "41"),
        ("4.3.3 System Monitoring, Customization, and Feedback Mechanisms", "42"),
        ("CHAPTER 5: RESEARCH FINDINGS AND DISCUSSIONS", "43"),
        ("CHAPTER 6: CONCLUSION AND RECOMMENDATIONS", "44"),
        ("REFERENCES", "46"),
        ("APPENDICES A: Geography", "47"),
        ("APPENDICES B: Data Structure and Frameworks", "48"),
        ("APPENDICES C: System Interface (UI)", "49"),
        ("APPENDICES D: Survey Questionnaires in Google Form", "54"),
        ("APPENDICES E: Activities And Planning", "58"),
        ("APPENDICES F: Photos", "63"),
    ]
    for title, page in contents:
        p = doc.add_paragraph()
        p.add_run(f"{title}\t{page}")
    add_page_break(doc)


def build_lists(doc: Document) -> None:
    add_centered(doc, "LIST OF TABLES", bold=True)
    doc.add_paragraph()
    tables = [
        "Table 1 List of participation",
        "Table 2 Validity Assessment of Research Instrument",
        "Table 3 Reliability Test Results (Cronbach's Alpha)",
        "Table 4 Effectiveness of Current Lost and Found Practices",
        "Table 5 Communication and Awareness Initiatives",
        "Table 6 Monitoring and Reporting Practices",
        "Table 7 Hypotheses and Appropriate SPSS Tests",
    ]
    for t in tables:
        doc.add_paragraph(t)
    add_page_break(doc)

    add_centered(doc, "LIST OF FIGURES", bold=True)
    doc.add_paragraph()
    figures = [
        "Figure 1 Conceptual framework for DLFS",
        "Figure 2 Map of Cambodia",
        "Figure 3 Dashboard DLFS",
        "Figure 4 Login Page",
        "Figure 5 Report Lost Item Page",
        "Figure 6 Report Found Item Page",
        "Figure 7 Item Detail and Claim Page",
        "Figure 8 Admin Claims Review Page",
        "Figure 9 Notifications Page",
        "Figure 10 Database Diagram",
        "Figure 11–30 Survey response charts (Questions 1–20)",
    ]
    for f in figures:
        doc.add_paragraph(f)
    add_page_break(doc)

    add_centered(doc, "LIST OF ABBREVIATIONS", bold=True)
    doc.add_paragraph()
    abbrevs = [
        ("DLFS", "Digital Lost & Found System"),
        ("DLFS", "BIU Lost and Found Management System"),
        ("BIU", "Beltei International University"),
        ("API", "Application Programming Interface"),
        ("REST", "Representational State Transfer"),
        ("ORM", "Object-Relational Mapping"),
        ("UI", "User Interface"),
        ("UX", "User Experience"),
        ("TAM", "Technology Acceptance Model"),
        ("RBV", "Resource-Based View"),
        ("JWT", "JSON Web Token"),
        ("SSE", "Server-Sent Events"),
        ("HTTP", "Hyper Text Transfer Protocol"),
        ("HTTPS", "Hyper Text Transfer Protocol Secure"),
        ("JSON", "JavaScript Object Notation"),
        ("SQL", "Structured Query Language"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("KPI", "Key Performance Indicator"),
        ("GDPR", "General Data Protection Regulation"),
    ]
    for short, full in abbrevs:
        doc.add_paragraph(f"{short}\t\t: {full}")
    add_page_break(doc)


def build_chapter1(doc: Document) -> None:
    add_heading(doc, "CHAPTER 1: INTRODUCTION")
    add_heading(doc, "1.1 Introduction to Research", level=2)
    add_body(
        doc,
        "In university environments, the loss of personal belongings is a common and stressful "
        "occurrence for students, faculty, and staff. Items such as student identification cards, "
        "mobile phones, wallets, keys, laptops, textbooks, USB drives, calculators, and headphones "
        "are frequently misplaced or lost on campus. Without an organized and accessible system to "
        "report and reclaim these items, the recovery process becomes frustrating, time-consuming, "
        "and often unsuccessful. For students in particular, losing essential belongings can disrupt "
        "daily attendance, delay access to classrooms and laboratories, prevent entry to examinations "
        "when ID cards are missing, and create unnecessary financial pressure when replacements must "
        "be purchased at short notice.",
    )
    add_body(
        doc,
        "Lost-and-found management is not merely an administrative task; it is a core student support "
        "function that reflects how well an institution cares for the daily wellbeing of its community. "
        "When students believe that the university provides reliable channels for recovering personal "
        "property, they develop greater trust in campus services. Conversely, when items remain lost "
        "despite being found by others, frustration grows and the perceived quality of student support "
        "declines. Effective lost-and-found coordination therefore contributes directly to student "
        "satisfaction, campus safety culture, and the overall reputation of the institution.",
    )
    add_body(
        doc,
        f"At {ORG} (BIU), the current approach to managing lost and found items relies on informal "
        "methods such as verbal announcements, social media posts, Telegram groups, Facebook pages, "
        "class group chats, and physical bulletin boards at the student center and front desk. These "
        "channels are inadequate and inconsistent because they lack structured search, centralized "
        "records, timestamped case history, standardized item categories, and a formal claim "
        "verification process. Information posted in one channel may never reach the owner who is "
        "checking a different platform. When important belongings are not recovered quickly, students "
        "experience academic disruption, financial loss, stress, and reduced confidence in campus "
        "support services. Staff members also spend additional time answering repeated inquiries, "
        "checking handwritten logs, and coordinating returns without a single authoritative system "
        "of record.",
    )
    add_body(
        doc,
        f"Common loss locations at {ORG} include classrooms, computer laboratories, the library, "
        "cafeteria areas, stairways, parking zones, and sports facilities. Items are most frequently "
        "misplaced during transitions between classes, when students move quickly and leave belongings "
        "on desks or benches. Found items may remain unclaimed for weeks because there is no central "
        "register linking finders to owners. In many cases, honest finders wish to return items but "
        "do not know the correct reporting procedure, while owners search multiple informal channels "
        "without success. This gap between intention and outcome demonstrates the urgent need for a "
        "structured digital solution.",
    )
    add_body(
        doc,
        "The evolution of technology has brought about new solutions to long-standing problems in "
        "educational institutions. Digital platforms, web applications, and database management "
        "systems have transformed administrative and student service functions across universities "
        "worldwide, including enrollment, library services, examination scheduling, fee payment, "
        "and campus announcements. A Digital Lost & Found System (DLFS) represents a practical and "
        "effective application of these technologies, enabling users to electronically report lost "
        "or found items, search through categorized listings, upload photos as evidence, receive "
        "automated match suggestions, and communicate securely to arrange item recovery (Johnson, "
        "2021). Such systems align with the broader trend of campus digital transformation, where "
        "services that were once handled manually are moved to transparent, searchable, and auditable "
        "online platforms.",
    )
    add_body(
        doc,
        "Compared with manual methods, a web-based lost-and-found platform offers several clear "
        "advantages. First, all reports are stored in one database rather than scattered across "
        "social media threads. Second, users can filter listings by item type, category, building, "
        "date, and keyword, dramatically reducing search time. Third, photo uploads improve "
        "identification accuracy and support fair claim verification. Fourth, administrators can "
        "track case status from open to resolved, measure recovery performance, and identify "
        "recurring problem areas on campus. Fifth, notifications alert users when potential matches "
        "appear or when claim decisions are made, eliminating the need to manually recheck multiple "
        "channels every day.",
    )
    add_body(
        doc,
        f"In Cambodia, higher education institutions are increasingly adopting information technology "
        f"to improve service delivery and meet the expectations of digitally connected students. "
        f"Young adults entering university today routinely use smartphones, social applications, and "
        f"online services in their daily lives. They reasonably expect campus administrative services "
        f"to offer similar convenience and reliability. As {ORG} expands its programs in Information "
        f"Technology, Business, Engineering, and Languages, the volume of daily campus activity—and "
        f"the number of personal items carried onto campus—continues to increase. A centralized "
        f"lost-and-found platform is therefore not only a convenience feature but also a necessary "
        f"component of modern campus infrastructure that supports student welfare, operational "
        f"efficiency, and institutional reputation.",
    )
    add_body(
        doc,
        f"This research explores the development and implementation of a web-based {SHORT_TITLE} "
        f"({SYSTEM_ACRONYM}) tailored specifically for {ORG}. The system is designed to address the "
        "inefficiencies of the current manual process by providing a centralized, accessible, and "
        "user-friendly platform. It allows students and staff to submit reports of lost or found "
        "items with descriptions and photos, enables keyword and category-based search functionality, "
        "supports match suggestions based on item metadata such as category, building, and event date, "
        "and facilitates claim submission with proof review by administrators (Miller, 2023). By "
        "replacing fragmented communication with a single trusted channel, DLFS reduces duplicate "
        "posts, improves accountability, and creates a clear audit trail from report to resolution.",
    )
    add_body(
        doc,
        f"The proposed {SYSTEM_ACRONYM} serves three primary stakeholder groups. For students and "
        "staff who lose items, the system provides a reliable place to report losses and monitor "
        "progress. For finders, it offers a simple and responsible way to submit found property "
        "without relying on informal posts that may be overlooked. For administrators and student "
        "affairs personnel, it delivers dashboards, claim review tools, and analytics that support "
        "efficient case management and evidence-based planning. Each group benefits from the same "
        "centralized data source, ensuring consistency and reducing miscommunication.",
    )
    add_body(
        doc,
        "The introduction of such a system is timely and necessary as BIU continues to grow in "
        "student population and campus size. A digital platform not only improves the likelihood of "
        "item recovery but also contributes to a more organized, transparent, and student-centric "
        "university environment. By automating item registration and tracking, the system reduces "
        "the burden on administrative staff while empowering users to take an active role in managing "
        "their lost belongings. Furthermore, it provides administrators with data insights on item "
        "trends, peak loss periods, category distribution, and recovery rates, supporting "
        "evidence-based decision-making for campus security and student affairs planning (Baker, 2023).",
    )
    add_body(
        doc,
        "From a software engineering perspective, the proposed system demonstrates how modern "
        "full-stack web development can solve a real campus problem using structured data models, "
        "secure authentication, validated forms, image storage, REST APIs, and role-based dashboards. "
        "The project applies industry-standard tools including Next.js, React, TypeScript, Prisma, "
        "PostgreSQL, NextAuth, Cloudinary, Zod validation, TanStack Table, Recharts, and OpenAPI "
        "documentation, ensuring that the solution is scalable, maintainable, and suitable for future "
        "extension such as mobile applications, Khmer language support, or SMS and email notifications. "
        "The development follows best practices in user experience design, including responsive layouts, "
        "accessible components, and real-time notification streams.",
    )
    add_body(
        doc,
        "Research into similar systems implemented at other institutions demonstrates their "
        "effectiveness in improving recovery rates and user satisfaction. Studies show that digital "
        "lost and found platforms with photo upload capabilities, real-time notifications, and "
        "claim verification mechanisms significantly outperform traditional paper-based approaches "
        f"(Nguyen, 2024). Survey research conducted as part of this study further confirms that BIU "
        f"students and staff strongly support adoption of a centralized platform and prefer features "
        f"such as photo evidence, searchable listings, and formal claim workflows. These findings "
        f"validate the direction of the {SYSTEM_ACRONYM} project and provide empirical support for "
        "campus deployment.",
    )
    add_body(
        doc,
        f"This research aims to analyze the specific needs of BIU through mixed-method investigation "
        f"including surveys, interviews, observation, and document analysis; design a system that "
        f"meets those needs; implement a production-ready web application; and evaluate its benefits "
        f"from the perspectives of students, staff, and administration. The study contributes both "
        f"practical value for {ORG} and academic value for the field of campus information systems "
        f"by demonstrating how a well-designed {SYSTEM_ACRONYM} can transform an overlooked manual "
        "process into an efficient digital service. The remainder of this chapter presents the research "
        "problem, objectives, research questions, significance, scope, limitations, and layout of "
        "the study.",
    )

    add_heading(doc, "1.2 Research Problem", level=2)
    add_body(
        doc,
        f"At {ORG}, the existing approach to managing lost and found items is unstructured and "
        "ineffective. Students and staff who lose items have limited reliable channels through which "
        "to report or search for their belongings. Similarly, those who find items have no standardized "
        "process for submitting them to a central repository. As a result, many items go unclaimed, "
        "and students who lose important belongings experience unnecessary stress and inconvenience.",
    )
    add_body(doc, "The major difficulties with the current situation at BIU are as follows:")
    add_bullets(
        doc,
        [
            "No centralized platform: information is scattered across social media, messaging groups, and front-desk notices.",
            "Poor searchability: users cannot filter by category, building, date, or keyword.",
            "Delayed communication: owners and finders often miss each other because updates are not tracked.",
            "No photo evidence standard: descriptions alone are insufficient for accurate identification.",
            "Weak claim verification: there is no formal workflow to prove ownership before returning items.",
            "Limited administrative visibility: staff cannot easily monitor open cases or measure recovery performance.",
            "Duplicate reporting: the same item may be posted multiple times in different channels.",
            "Low recovery rates: many lost items are never reunited with their owners.",
        ],
    )
    add_body(
        doc,
        f"Therefore, this research addresses the lack of a comprehensive, integrated {SYSTEM_ACRONYM} "
        f"that can manage the full lost-and-found workflow at {ORG}, from reporting through claiming "
        "and resolution.",
    )

    add_heading(doc, "1.3 Research Aim/Objective", level=2)
    add_body(
        doc,
        f"The main goal of this research is to design, implement, and evaluate a {SYSTEM_ACRONYM} "
        f"that improves the speed, accuracy, and transparency of campus lost-and-found management "
        f"at {ORG}. The specific objectives include:",
    )
    add_bullets(
        doc,
        [
            f"To identify the key components and functionalities required for a successful {SYSTEM_ACRONYM}.",
            "To examine the impact of the digital system on item recovery efficiency and claim traceability.",
            "To assess the system's role in improving communication between students, finders, and staff.",
            f"To analyze the advantages and challenges of {SYSTEM_ACRONYM} compared to manual lost-and-found methods.",
            "To conduct a literature review on campus information systems and digital service management.",
            "To identify key challenges faced by users in reporting, searching, and claiming items.",
            "To evaluate user readiness for adopting a centralized lost-and-found web platform.",
            "To provide recommendations for sustainable deployment and future enhancement.",
        ],
    )

    add_heading(doc, "1.3.1 Research Question", level=2)
    add_body(
        doc,
        "The key research issue for this study is: how can a Digital Lost & Found System effectively "
        "address the challenges of item recovery at BELTEI International University, and what features "
        "are essential for maximizing its benefits for students, staff, and administration? "
        "To address this main question, the following specific research questions will be investigated:",
    )
    add_bullets(
        doc,
        [
            "What are the benefits of this system for BELTEI International University and its students?",
            "Is the system easy to use for reporting and searching lost or found items?",
            "How does the system work and provide a convenient process for claiming returned items?",
            "How does the system improve communication between item finders and owners?",
            "What improvements can be made to increase the system's effectiveness and user satisfaction?",
        ],
    )

    add_heading(doc, "1.3.2 Signification of the study", level=2)
    add_body(
        doc,
        "The significance of this study lies in its contribution to the digital transformation of "
        "campus support services at BIU. It provides a practical model for universities seeking to "
        "replace fragmented lost-and-found practices with a secure, searchable, and auditable platform. "
        "The findings are valuable for software developers, campus administrators, student affairs "
        "teams, and IT departments responsible for service innovation.",
    )
    add_body(
        doc,
        "For students, the system reduces the time and stress associated with losing personal items "
        "by providing a reliable channel for reporting and searching. For staff and administrators, "
        "the system improves operational efficiency, claim accountability, and service transparency. "
        "For the institution, successful deployment enhances BIU's reputation as a modern university "
        "that invests in practical digital solutions for everyday campus needs.",
    )

    add_heading(doc, "1.4 Scope and limitation", level=2)
    add_body(
        doc,
        f"The study focuses on the analysis, design, implementation, and evaluation of {SYSTEM_ACRONYM} "
        f"at {ORG}. The scope covers lost item reporting, found item reporting, image uploads via "
        "Cloudinary, search and filtering, match suggestions, claim submission and review, in-app "
        "notifications, user dashboards, admin monitoring, and REST API documentation.",
    )

    add_heading(doc, "1.4.1 Scope and Limitations", level=2)
    add_body(
        doc,
        f"This study examines lost-and-found management within the {ORG} campus context. Data were "
        f"collected from {SAMPLE_SIZE} respondents through a structured Google Form questionnaire, "
        "supplemented by interviews and observation. The system is implemented as a production-ready "
        "web application using Next.js 16, React 19, TypeScript, Prisma ORM, PostgreSQL (Neon), "
        "NextAuth v5, Cloudinary, Zod validation, React Hook Form, TanStack Table, Recharts, and "
        "OpenAPI/Swagger documentation.",
    )
    add_body(
        doc,
        "The research emphasizes improvements in reporting quality, searchability, claim verification, "
        "notification timeliness, and administrative analytics. The study does not cover hardware "
        "integration such as RFID lockers or biometric access systems, although these may be considered "
        "in future extensions.",
    )

    add_heading(doc, "1.4.2 Limitations of the Study", level=2)
    add_body(
        doc,
        f"Despite careful planning, this study has several limitations. First, the research was "
        f"conducted only at {ORG}, which may limit generalization to other universities with different "
        f"campus sizes or IT infrastructures. Second, although {SAMPLE_SIZE} survey responses were "
        "collected, some respondents may not have personally experienced item loss, which could influence "
        "perception-based answers. Third, the study relies partly on self-reported survey data, which "
        "may be subject to response bias. Fourth, time constraints limited long-term measurement of "
        "actual recovery rates after full campus-wide deployment. Finally, privacy requirements "
        "restricted access to some operational records during analysis.",
    )

    add_heading(doc, "1.5 Layout of the study/Research", level=2)
    add_body(
        doc,
        f"CHAPTER 1 INTRODUCTION, introduces the background of the study by highlighting the "
        f"importance of effective lost-and-found management at {ORG}. The chapter focuses on "
        f"identifying the research problem, objectives, research questions, significance, scope, "
        f"and limitations related to the current manual lost-and-found practices. Key issues such "
        f"as fragmented reporting channels, poor searchability, delayed communication, and weak "
        f"claim verification are examined, while the research aim of designing a centralized "
        f"digital platform is clearly defined. The chapter also establishes the foundation for the "
        f"entire study and guides the direction of subsequent chapters.",
    )
    add_body(
        doc,
        f"CHAPTER 2 LITERATURE REVIEW, reviews relevant literature related to lost-and-found "
        f"systems, campus information services, and digital platform adoption in educational "
        f"institutions. The review focuses on identifying key definitions, theories, and models "
        f"that support the development of the proposed system. Concepts such as the Technology "
        f"Acceptance Model (TAM), Service Quality Theory (SERVQUAL), and the Resource-Based View "
        f"(RBV) are examined, while prior studies on usability, notification systems, privacy "
        f"protection, and claim verification are highlighted. The chapter also presents the "
        f"conceptual framework that explains the relationship between system features, "
        f"{SYSTEM_ACRONYM}, and improved campus service outcomes.",
    )
    add_body(
        doc,
        f"CHAPTER 3 RESEARCH METHODOLOGY, describes the research methodology used to conduct the "
        f"study at {ORG}. The chapter focuses on explaining the research design, research "
        f"location, population, sampling technique, and data collection instruments. Methods such "
        f"as a Google Form questionnaire distributed to {SAMPLE_SIZE} respondents, semi-structured "
        f"interviews, direct observations, and document analysis are examined, while the UI and "
        f"UX scope of the proposed system is clearly described. The chapter also explains the "
        f"technical architecture of the Digital Lost & Found System ({SYSTEM_ACRONYM}), including "
        f"Next.js, Prisma, PostgreSQL, NextAuth, Cloudinary, and OpenAPI documentation.",
    )
    add_body(
        doc,
        f"CHAPTER 4 DATA ANALYSIS, presents the analysis of data collected from respondents at "
        f"{ORG}. The analysis focuses on identifying the strengths and weaknesses of the current "
        f"lost-and-found management practices. Strengths such as staff willingness to help, "
        f"informal peer-to-peer communication, and existing front-desk support are examined, while "
        f"weaknesses such as manual processes, delayed response, lack of photo evidence, scattered "
        f"information, and absence of claim tracking are highlighted. The chapter also proposes "
        f"suitable solutions to address these weaknesses through the implementation of a Digital "
        f"Lost & Found System ({SYSTEM_ACRONYM}), including web, REST API, and mobile integration.",
    )
    add_body(
        doc,
        f"CHAPTER 5 RESEARCH FINDINGS AND DISCUSSIONS, presents the key research findings derived "
        f"from the survey and supporting qualitative data collected at {ORG}. The chapter focuses "
        f"on interpreting user perceptions, current practices, system usefulness, privacy "
        f"concerns, feature preferences, and readiness for digital adoption. Findings such as high "
        f"demand for photo uploads, search filters, and claim verification are examined, while "
        f"concerns regarding data privacy and platform trust are highlighted. The chapter also "
        f"discusses the implications of the findings in relation to the research objectives and "
        f"the successful deployment of {SYSTEM_ACRONYM}.",
    )
    add_body(
        doc,
        f"CHAPTER 6 CONCLUSION AND RECOMMENDATIONS, concludes the study by summarizing the overall "
        f"findings and addressing the research objectives and questions. The chapter focuses on "
        f"confirming the need to transition from manual and informal lost-and-found coordination "
        f"to a centralized digital system at {ORG}. Conclusions such as improved recovery "
        f"efficiency, better communication, and stronger administrative control are examined, while "
        f"practical recommendations for campus adoption, staff training, and future system "
        f"enhancement are highlighted. The chapter also provides guidance for sustainable "
        f"implementation of the Digital Lost & Found System ({SYSTEM_ACRONYM}).",
    )
    add_body(
        doc,
        "In addition to the six chapters, the thesis includes REFERENCES and APPENDICES. Appendix A "
        "presents geographic context. Appendix B describes the database structure and conceptual "
        "framework. Appendix C provides system interface screenshots. Appendix D contains the full "
        f"Google Form survey questionnaire. Appendix E outlines the project activities and planning "
        "timeline. Appendix F includes development progress photographs.",
    )
    add_page_break(doc)


def build_chapter2(doc: Document) -> None:
    add_heading(doc, "CHAPTER 2: LITERATURE REVIEW")
    add_heading(doc, "2.1 Definitions and Theories Related to the Topic", level=2)
    add_body(
        doc,
        "The concept of a Digital Lost & Found System draws from multiple disciplines, including "
        "information systems design, human-computer interaction, and campus service management. "
        "Several key definitions and theoretical frameworks are important for understanding the system, "
        "how it is designed, and its impact on campus life. This section establishes the conceptual "
        f"foundation for the {SYSTEM_ACRONYM} research by clarifying essential terms and explaining "
        "the theories that guide system design, user adoption, and service quality evaluation.",
    )
    add_body(
        doc,
        "Lost & Found System: A Lost & Found System refers to an organized process or platform used "
        "by an institution to receive, catalog, store, and reunite lost items with their rightful "
        "owners. Traditional systems rely on physical offices, notice boards, and manual record "
        "keeping. Digital versions leverage web or mobile platforms to extend the reach and efficiency "
        "of physical lost and found offices by enabling online reporting, searchable databases, and "
        "automated notifications (Cohen & Swerdlik, 2018). At BIU, transitioning to a digital "
        "lost-and-found system means replacing fragmented informal channels with a structured service "
        "that supports the full lifecycle of a lost or found item.",
    )
    add_body(
        doc,
        "Information System: An information system is a combination of people, processes, data, and "
        "technology used to collect, store, process, and distribute information to support "
        "organizational decision-making and operations. In the context of campus services, information "
        "systems such as DLFS help administrators manage records accurately while giving users timely "
        "access to the information they need to recover personal property (Laudon & Laudon, 2020).",
    )
    add_body(
        doc,
        "Web-Based Platform: A web-based platform is an application accessible via a web browser "
        "over the internet or an intranet. Web-based platforms are preferred in educational contexts "
        "due to their accessibility across different devices and operating systems without requiring "
        "additional software installation. Students can access the system from laptops, tablets, or "
        "smartphones, making it suitable for a mobile student population (Anderson, 2008).",
    )
    add_body(
        doc,
        "User Interface (UI) and User Experience (UX): The user interface refers to the visual layout "
        "and interactive elements through which users operate a system, including forms, buttons, menus, "
        "and navigation. User experience encompasses all aspects of a user's interaction with a digital "
        "product, including usability, accessibility, efficiency, and satisfaction. A positive UX is "
        "critical for the adoption and effective use of digital campus services because students are "
        "more likely to report and search for items when the process is simple and trustworthy "
        "(Wiggins, 1990; Nielsen, 2012).",
    )
    add_body(
        doc,
        "Item Recovery: Item recovery refers to the successful return of a lost item to its legitimate "
        "owner. Recovery rate is a key performance indicator for lost-and-found services and reflects "
        "how effectively the institution connects finders with owners. Factors that improve recovery "
        "include accurate item descriptions, photo evidence, timely notifications, searchable records, "
        "and formal claim verification procedures.",
    )
    add_body(
        doc,
        "Claim Verification: Claim verification is the process of confirming that a person claiming "
        "a found item is its legitimate owner. This may involve providing item descriptions, photos, "
        "identification, proof of purchase, or answering specific questions about the item. Digital "
        "systems support claim verification through structured claim forms, admin review workflows, "
        "and audit trails that reduce disputes and prevent fraudulent claims (Cronbach, 1970).",
    )
    add_body(
        doc,
        "Digital Transformation: Digital transformation is the integration of digital technology into "
        "all areas of an organization, fundamentally changing how services are delivered and how value "
        "is created for users. In higher education, digital transformation includes moving manual "
        "campus services—such as lost-and-found coordination—to online platforms that are faster, more "
        "transparent, and easier to manage at scale.",
    )
    add_body(
        doc,
        "Technology Acceptance Model (TAM): The Technology Acceptance Model posits that perceived "
        "usefulness and perceived ease of use are primary determinants of technology adoption. Perceived "
        "usefulness refers to the degree to which a person believes that using a system will improve "
        "performance, such as recovering lost items more quickly. Perceived ease of use refers to the "
        f"degree to which a person believes that using the system requires minimal effort. TAM is "
        f"directly relevant to assessing whether BIU students and staff will adopt {SYSTEM_ACRONYM} "
        "(Davis, 1989).",
    )
    add_body(
        doc,
        "Service Quality Theory (SERVQUAL): SERVQUAL provides a framework for evaluating service quality "
        "based on five dimensions: reliability, responsiveness, assurance, empathy, and tangibles. "
        "Applied to DLFS, reliability means consistent and accurate item records; responsiveness means "
        "timely notifications and claim decisions; assurance means secure handling of personal data; "
        "empathy means understanding the stress users feel when items are lost; and tangibles means a "
        "professional and easy-to-use interface (Parasuraman, Zeithaml, & Berry, 1988).",
    )
    add_body(
        doc,
        "Resource-Based View (RBV): The Resource-Based View explains how organizations gain competitive "
        "advantage by developing valuable, rare, and difficult-to-imitate resources. A well-implemented "
        "digital lost-and-found platform becomes a strategic institutional resource that improves "
        "operational efficiency, strengthens student support services, and enhances the university's "
        "reputation for modern campus management (Barney, 1991).",
    )
    add_body(
        doc,
        f"Together, these definitions and theories provide the conceptual basis for designing, "
        f"implementing, and evaluating {SYSTEM_ACRONYM}. They explain why users adopt digital platforms, "
        f"how service quality should be measured, and how the system contributes to institutional "
        f"performance at {ORG}.",
    )

    add_heading(doc, "2.2 Analysis of Scholars' Concepts/Words/Sayings", level=2)
    add_body(doc, "Scholars have offered diverse perspectives on digital service systems in educational institutions:")
    add_bullets(
        doc,
        [
            "Usability and Accessibility: Davis (1989) demonstrated that perceived ease of use is a strong predictor of technology acceptance, particularly among student populations with varying digital literacy.",
            "Digital Service Integration: Institutions that integrate digital platforms into campus operations report improved student satisfaction and administrative efficiency (Creswell & Plano, 2017).",
            "Real-Time Notification Systems: Instant alerts when matching items are reported significantly improve recovery speed and user engagement (Black & Wiliam, 1998).",
            "Security and Privacy: Users must trust that personal information is protected during reporting and claiming (McCabe, 2001).",
            "Photo Documentation: Visual evidence substantially improves item identification and reduces fraudulent claims (Thompson, 2022).",
        ],
    )

    add_heading(doc, "2.3 Experiences or Issues Raised", level=2)
    add_bullets(
        doc,
        [
            "Low Awareness and Adoption: Without proper marketing and orientation, even well-designed systems may go underutilized (Smith, 2022).",
            "Technical Challenges: Downtime, slow loading, and device compatibility issues reduce trust in the platform (Johnson, 2021).",
            "Data Management: Growing databases require archiving, retention policies, and regular maintenance (Elmasri & Navathe, 2016).",
            "Cultural Considerations: Some users may hesitate to report found items through digital channels; understanding local context is essential at BIU.",
            "Claim Fraud: Without robust verification, systems are vulnerable to fraudulent claims; photo and admin review mitigate this risk (Baker, 2023).",
        ],
    )

    add_heading(doc, "2.4 Research Model or Framework", level=2)
    add_body(
        doc,
        f"The research model explains how the implementation of {SYSTEM_ACRONYM} can improve lost-and-found "
        f"management at {ORG}. As shown in Figure 1, the framework contains independent variables, an "
        "intervening variable, and dependent variables that describe the cause-and-effect relationship "
        "between system features, platform implementation, and campus service outcomes.",
    )
    add_body(
        doc,
        "The independent variables represent the core functional features of the proposed system: system "
        "usability, system accessibility, reporting efficiency, search functionality, notification systems, "
        f"and claim verification. These inputs are processed through {SYSTEM_ACRONYM} as the intervening "
        "variable, which provides a centralized item database, real-time notifications, data security and "
        "privacy controls, and user acceptance support. The dependent variables measure the expected "
        "outcomes: increased recovery rates, improved finder-owner communication, reduced search time, "
        "higher user satisfaction, enhanced administrative efficiency, and a more organized campus "
        "environment.",
    )
    add_figure(doc, FIGURE_1, "Figure 1. Conceptual framework for DLFS")
    add_page_break(doc)


CHAPTER3_SURVEY_SECTIONS: list[tuple[str, list[tuple[str, str]]]] = [
    (
        "Section A — Respondent Profile (Questions 1–3)",
        [
            (
                "QUESTION 1: What is your campus role?",
                "The distribution shows that 64.3% (148) of respondents are students, 18.7% (43) are "
                "academic staff, 9.1% (21) are administrative staff, and 7.8% (18) are security or "
                "front-desk staff. This confirms that the sample is student-majority while still "
                "including operational stakeholders who manage lost-and-found intake and verification.",
            ),
            (
                "QUESTION 2: What is your gender?",
                "The chart shows 62.2% male (143), 36.1% female (83), and 1.7% prefer not to say (4) "
                "across 230 respondents. The sample includes both male and female campus users, "
                "providing a balanced perspective on lost-and-found service needs.",
            ),
            (
                "QUESTION 3: What is your approximate age?",
                "Most respondents (68.7%, 158) are aged 18–22, followed by 23–27 (20.4%, 47), under 18 "
                "(6.5%, 15), and 28 and above (4.3%, 10). The age distribution reflects BIU's young "
                "student population and suggests that survey findings are highly relevant to the main "
                "campus user group.",
            ),
        ],
    ),
    (
        "Section B — Lost-and-Found Experience (Questions 4–6)",
        [
            (
                "QUESTION 4: How often have you lost an item on campus in the past year?",
                "Results show that 38.7% (89) lost items two to three times, 26.5% (60) lost items "
                "more than three times, 26.1% (61) lost items once, and only 8.7% (20) reported no loss "
                "in the past year. In total, 91.3% of respondents experienced at least one lost-item "
                "incident, confirming that lost-item issues are common and that DLFS addresses a real "
                "campus problem.",
            ),
            (
                "QUESTION 5: What is your faculty?",
                "Respondents came from multiple faculties: Information Technology and Science (46.1%, 106), "
                "Business and Economics (19.6%, 45), Education (10.9%, 25), Law and Social Sciences "
                "(9.1%, 21), Engineering (7.8%, 18), and Other (6.5%, 15). This diversity strengthens "
                "the generalizability of findings across BIU programs.",
            ),
            (
                "QUESTION 6: What type of item did you most recently lose or find?",
                "Documents or ID cards accounted for 21.7% (50) of responses, followed by electronics "
                "(18.3%, 42), clothing (17.4%, 40), bags or luggage (14.8%, 34), keys or access cards "
                "(14.3%, 33), and other items (13.5%, 31). These categories align with the item "
                "classification used in the DLFS database schema.",
            ),
        ],
    ),
    (
        "Section C — Platform Perceptions (Questions 7–11)",
        [
            (
                "QUESTION 7: What method did you use to search for or report your lost item?",
                "Social media was used by 47.4% (109) of respondents, the security office by 22.6% (52), "
                "no action taken by 17.0% (39), and friends or classmates by 13.0% (30). This confirms "
                "fragmented reporting channels and supports the need for a centralized digital platform.",
            ),
            (
                "QUESTION 8: Were you able to recover your lost item?",
                "Only 34.8% (80) answered Yes, while 65.2% (150) answered No. This recovery gap "
                "demonstrates a major weakness in current practices and supports the implementation of "
                "searchable listings and match suggestions in DLFS.",
            ),
            (
                "QUESTION 9: How useful would a Digital Lost & Found platform be?",
                "On a 1–5 scale, 62.2% rated the platform 4 or 5 (useful to extremely useful), 21.7% "
                "rated it neutral (3), and 16.1% rated it 1 or 2. The mean score was 3.67 out of 5. "
                "Strong perceived usefulness supports the Technology Acceptance Model expectation that "
                "users will adopt DLFS when they believe it improves item recovery.",
            ),
            (
                "QUESTION 10: How important is privacy protection when reporting lost or found items?",
                "Privacy was rated 4 or 5 (important to extremely important) by 75.2% of respondents, "
                "with 46.5% selecting the highest rating. The mean score was 4.04 out of 5. This finding "
                "supports the inclusion of role-based access, secure authentication, and controlled "
                "contact visibility in the system design.",
            ),
            (
                "QUESTION 11: Would you use a digital platform to report found items?",
                "A total of 84.8% (195) answered Yes, 3.9% (9) answered Maybe, and 11.3% (26) answered No. "
                "High willingness to report found items indicates strong readiness for campus-wide DLFS "
                "adoption.",
            ),
        ],
    ),
    (
        "Section D — System Feature Evaluation (Questions 12–14)",
        [
            (
                "QUESTION 12: Would uploading photos of items help verification?",
                "A total of 77.8% (179) answered Yes, 7.8% (18) answered No, and 14.3% (33) were Not sure. "
                "Strong support for photo upload confirms the need for image attachments in lost and "
                "found reports, consistent with Photo Upload ranking first among preferred features in "
                "Question 15.",
            ),
            (
                "QUESTION 13: Would searchable online listings with filters (category, building, date) help?",
                "A total of 81.3% (187) answered Yes, 6.1% (14) answered No, and 12.6% (29) were Not sure. "
                "This validates the inclusion of keyword search, category filters, building filters, and "
                "date sorting in the DLFS browse page.",
            ),
            (
                "QUESTION 14: Would a digital claim workflow (submit proof → staff review → notification) be useful?",
                "A total of 79.6% (183) answered Yes, 7.0% (16) answered No, and 13.4% (31) were Not sure. "
                "Respondents expect a structured claim process with proof submission and admin review, "
                "which directly informed the claims module and notification system in DLFS.",
            ),
        ],
    ),
    (
        "Section E — Feature Preferences (Question 15)",
        [
            (
                "QUESTION 15: Which features are most important to you?",
                "Respondents selected multiple preferred features: Photo Upload (60.9%, 140), Keyword Search "
                "(57.4%, 132), Claim Verification (57.0%, 131), Location Filtering (55.2%, 127), and "
                "Real-time Notifications (53.0%, 122). These results directly informed the functional "
                "requirements of DLFS.",
            ),
        ],
    ),
    (
        "Section F — Motion Graphics Evaluation (Questions 16–20)",
        [
            (
                "QUESTION 16: Motion graphics help navigation",
                "The mean agreement score was 4.13 out of 5, with 46.5% strongly agreeing. Respondents "
                "value animated UI guidance when using campus web applications.",
            ),
            (
                "QUESTION 17: Animated transitions improve the user experience",
                "The mean agreement score was 4.05 out of 5, with 43.5% strongly agreeing. This supports "
                "the use of smooth page transitions in the DLFS interface.",
            ),
            (
                "QUESTION 18: Loading animations reduce frustration during data retrieval",
                "The mean agreement score was 4.04 out of 5, with 44.3% strongly agreeing. Skeleton "
                "loaders and progress indicators improve perceived system responsiveness.",
            ),
            (
                "QUESTION 19: Motion-based icons improve usability",
                "The mean agreement score was 4.17 out of 5, with 45.7% strongly agreeing. This supports "
                "icon animation for key actions such as report, search, and claim.",
            ),
            (
                "QUESTION 20: Motion graphics improve overall system quality",
                "The mean agreement score was 4.18 out of 5, the highest among motion graphics items, "
                "with 45.7% strongly agreeing. This supports investment in polished UI motion as part "
                "of overall service quality.",
            ),
        ],
    ),
]


def build_chapter3(doc: Document) -> None:
    add_heading(doc, "CHAPTER 3: RESEARCH METHODOLOGY")
    add_body(
        doc,
        f"This chapter describes the methodology used at {ORG} to carry out the study on campus "
        f"lost-and-found management. It explains the population, setting, data gathering "
        f"strategies, sampling tactics, and research methodology. The approach guarantees that "
        f"the results are precise, trustworthy, and representative of the campus lost-and-found "
        f"service practices.",
    )
    add_heading(doc, "3.1 Research Methodology", level=2)
    add_body(
        doc,
        f"This study used a mixed-method approach at {ORG} to understand real-world lost-and-found "
        f"challenges and evaluate the need for {SYSTEM_ACRONYM}. Qualitative methods included "
        f"interviews and observation. Quantitative methods included structured questionnaires "
        f"distributed to students and staff.",
    )
    add_body(
        doc,
        f"This study's research technique was thoughtfully created to completely comprehend the "
        f"difficulties in real-world campus lost-and-found coordination. A mixed-method strategy "
        f"that included qualitative and quantitative techniques was used to do this, guaranteeing "
        f"that both quantifiable data and in-depth insights were gathered and examined. The "
        f"qualitative component concentrated on using semi-structured interviews and direct "
        f"observation to collect in-depth viewpoints from students, front-desk staff, and "
        f"administrators. The quantitative component entailed delivering organized questionnaires "
        f"to campus users with scaled and closed-ended questions for statistical analysis. By "
        f"integrating these methods, the study created a strong basis for the Digital Lost & Found "
        f"System ({SYSTEM_ACRONYM}) design, guaranteeing that the suggested remedy directly "
        f"addressed the issues and requirements noted in the research.",
    )

    add_heading(doc, "3.1.1 Research Design", level=2)
    add_body(
        doc,
        "A descriptive and analytical research design was adopted. The descriptive component "
        "documented current practices and user experiences. The analytical component evaluated "
        "weaknesses and proposed an automated web-based solution.",
    )
    add_body(
        doc,
        f"This study adopts a descriptive and analytical research design that integrates both "
        f"quantitative and qualitative research methods to comprehensively examine the existing "
        f"lost-and-found management practices at {ORG}. The descriptive approach is used to "
        f"clearly identify and document the current practices, procedures, and challenges "
        f"associated with reporting, searching, and recovering lost items, while the analytical "
        f"approach allows for a deeper evaluation of the effectiveness, accuracy, and efficiency "
        f"of these practices.",
    )
    add_body(
        doc,
        f"The primary objective of this research design is to identify existing challenges in "
        f"manual lost-and-found coordination, assess their impact on campus service performance, "
        f"and propose an automated solution that can enhance accuracy, efficiency, and "
        f"transparency in item recovery management. Quantitative data were collected through "
        f"structured surveys distributed to students and staff. Qualitative data were gathered "
        f"through semi-structured interviews with front-desk staff and administrators, direct "
        f"observations of reporting workflows, and document analysis of notice-board records and "
        f"informal social media posts.",
    )

    add_heading(doc, "3.1.2 Research Location", level=2)
    add_body(
        doc,
        f"The research was conducted at {ORG} in Phnom Penh, Cambodia. Phnom Penh is the capital "
        f"and a major center for higher education and technology adoption, making it a suitable "
        f"environment for studying campus digital services.",
    )
    add_body(
        doc,
        f"The research took place in Phnom Penh, the capital of Cambodia, encompassing 678.46 "
        f"square kilometers across 105 Sangkats and 14 Khan; nevertheless, this research "
        f"pertains solely to the {ORG} campus. The university was selected because its "
        f"lost-and-found coordination still depends on conventional methods such as front-desk "
        f"notices, informal messaging groups, and scattered social media posts. The study is "
        f"applicable to similar universities because it represents a typical higher-education "
        f"campus environment in Cambodia.",
    )
    add_centered(doc, "Figure 2. Map of Cambodia")

    add_heading(doc, "3.1.3 Research Population", level=2)
    add_body(
        doc,
        "The sample was chosen using purposive sampling, ensuring that all key stakeholders in "
        "campus lost-and-found management were represented.",
    )
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Campus Role", "Responses", "Percentage"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ("Student", "148", "64.3%"),
        ("Academic Staff", "43", "18.7%"),
        ("Administrative Staff", "21", "9.1%"),
        ("Security / Front Desk Staff", "18", "7.8%"),
    ]
    for r, row in enumerate(data, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val
    table.rows[5].cells[0].text = "Total"
    table.rows[5].cells[1].text = str(SAMPLE_SIZE)
    table.rows[5].cells[2].text = "100%"
    doc.add_paragraph("Table 1 List of participation")

    add_heading(doc, "3.1.4 Research Data Scope of Database", level=2)
    add_body(doc, "Database")
    add_body(
        doc,
        "PostgreSQL is an open-source relational database management system. Its name combines "
        '"Post" from the original project at the University of California, Berkeley, and "SQL", '
        "the acronym for Structured Query Language. For the Digital Lost & Found System (DLFS), "
        "PostgreSQL was selected as the database engine for several strategic reasons. Its "
        "open-source nature makes it a cost-effective solution, which is important for academic "
        "and institutional projects with budgetary considerations. PostgreSQL's widespread "
        "support and compatibility across platforms and programming languages enhance usability "
        "and integration with Next.js and Prisma ORM. It is recognized for high performance when "
        "managing large volumes of structured data, which is essential for lost-and-found listings, "
        "claims, and notifications. PostgreSQL also offers robust security features—including "
        "role-based access and encrypted connections—which safeguard sensitive user and contact "
        "data. Its scalability ensures the database can grow with campus adoption, supporting "
        "both pilot deployment and full institutional use. A large PostgreSQL community provides "
        "resources for support and knowledge sharing. Combined with proven reliability and ease "
        "of management through Prisma ORM and Neon cloud hosting, PostgreSQL is an ideal choice "
        "for a dependable, efficient, and versatile database engine for DLFS (PostgreSQL Global "
        "Development Group, 1996).",
    )
    add_body(
        doc,
        'Database Schema: The database schema is the structure of a database described in a '
        "formal language supported typically by a relational database management system. The term "
        '"schema" refers to the organization of data as a blueprint of how the database is '
        "constructed.",
    )
    add_body(
        doc,
        "Database schema design provides a blueprint to develop the architecture of the database, "
        "so a large amount of information can be stored systematically. It also refers to the "
        "strategies and best practices involved in constructing a database. Database schema design "
        "makes data much easier to consume, interpret, and retrieve by organizing it into "
        "separate entities and determining the relationships between organized entities.",
    )
    add_body(
        doc,
        f"For {SYSTEM_ACRONYM}, the schema was implemented using Prisma ORM and PostgreSQL (Neon). "
        "Core entities include User, Item, Claim, Notification, Account, Session, and "
        "VerificationToken.",
    )
    add_body(doc, "Item records store:")
    add_bullets(
        doc,
        [
            "type (LOST / FOUND)",
            "status (OPEN / RESOLVED / CLOSED)",
            "category (ELECTRONICS, DOCUMENTS, KEYS, CLOTHING, BOOKS, ACCESSORIES, SPORTS, STATIONERY, BAGS, OTHER)",
            "building and room hint",
            "title, description, color, brand",
            "event date and approximate time",
            "image URLs and contact preferences",
            "found disposition (STILL_HAVE, SUBMITTED_SECURITY, LEFT_WHERE_FOUND)",
        ],
    )
    add_body(doc, "Claim records store:")
    add_bullets(
        doc,
        [
            "claim type (FINDER / OWNER)",
            "proof message and proof image URLs",
            "review status (PENDING / APPROVED / REJECTED)",
            "admin note and reviewed timestamp",
        ],
    )
    add_body(doc, "Notification records store:")
    add_bullets(
        doc,
        [
            "kind (SYSTEM, MATCH, CLAIM, ITEM)",
            "title, message, link",
            "read status and creation time",
        ],
    )
    add_body(
        doc,
        "Relationships: one User may post many Items; one User may submit many Claims; one Item "
        "may have many Claims; one User receives many Notifications. Account and Session tables "
        "support NextAuth authentication.",
    )
    add_centered(doc, "Figure 10 Database Diagram")

    add_heading(doc, "3.1.5 Research Data Scope of UI", level=2)
    add_body(
        doc,
        f"The User Interface (UI) Scope of the proposed Digital Lost & Found System ({SYSTEM_ACRONYM}) "
        f"focuses on creating a simple and user-friendly interface to improve item reporting, "
        f"searching, claiming, and administration. Key UI design considerations include clarity, "
        f"mobile responsiveness, accessibility, and consistent navigation across student and admin "
        f"workflows.",
    )
    add_body(doc, "Key UI pages implemented in DLFS include:")
    add_bullets(
        doc,
        [
            "Home page with recent lost and found listings",
            "Login and registration pages",
            "Multi-step Report Lost and Report Found forms",
            "Browse and search page with filters",
            "Item detail page with claim action",
            "User dashboard with statistics and charts",
            "Notifications page with live updates",
            "Admin dashboard and claims review panel",
            "API documentation page (/api-docs)",
        ],
    )
    add_centered(doc, "Figure 3 Dashboard DLFS")
    add_body(doc, "Backend and platform technologies:")
    add_body(
        doc,
        "Next.js App Router: Next.js is a full-stack React framework that supports server "
        "components, App Router navigation, and API routes in a single codebase. This allows "
        f"{SYSTEM_ACRONYM} to render pages efficiently on the server while still providing "
        "interactive client components where needed.",
    )
    add_body(
        doc,
        "Prisma ORM: Prisma is a type-safe Object-Relational Mapping tool used to access the "
        "PostgreSQL database defined in Section 3.1.4. It defines models for User, Item, Claim, "
        "Notification, Account, and Session, ensuring consistent data structure and compile-time "
        "type checking.",
    )
    add_body(
        doc,
        "NextAuth v5: NextAuth provides secure authentication for the web application. It "
        "supports email/password credentials and optional social login providers. Session "
        "management ensures that only authenticated users can manage listings, submit claims, "
        "and access personal dashboards.",
    )
    add_body(
        doc,
        "Cloudinary: Cloudinary is used for image upload, storage, optimization, and delivery. "
        "Students can attach up to five photos per lost or found report, and claimants can upload "
        "proof images. This improves verification quality and reduces disputes over item identity.",
    )
    add_body(
        doc,
        "Zod + React Hook Form: Zod provides TypeScript-first schema validation for item report "
        "forms, claim forms, and profile forms. React Hook Form manages multi-step lost and found "
        "reporting with real-time validation feedback.",
    )
    add_body(
        doc,
        "OpenAPI + Swagger UI: The REST API is documented using OpenAPI and exposed through "
        "Swagger UI at /api-docs. This supports integration with mobile apps and external clients.",
    )
    add_body(
        doc,
        "REST API: Web services follow REST architecture. Resources include items, claims, "
        "notifications, uploads, profile, and authentication. JSON is used for request and "
        "response bodies.",
    )
    add_centered(doc, "Figure 4 Login Page")
    add_centered(doc, "Figure 5 Report Lost Item Page")
    add_centered(doc, "Figure 6 Report Found Item Page")
    add_centered(doc, "Figure 7 Item Detail and Claim Page")
    add_centered(doc, "Figure 8 Admin Claims Review Page")
    add_body(doc, "Frontend technologies:")
    add_bullets(
        doc,
        [
            "React 19 + TypeScript strict mode: React powers the interactive user interface with strict TypeScript typing to prevent runtime errors and improve maintainability.",
            "Tailwind CSS and shadcn/ui component patterns: Tailwind CSS provides utility-first styling with semantic design tokens for light and dark themes.",
            "TanStack Table for admin data views: Admin pages use TanStack Table to display users, items, and claims with sorting, filtering, and pagination.",
            "Recharts for dashboard analytics: Dashboard charts visualize items over time, category distribution, and resolution rates.",
            "Server-Sent Events for live notifications: Real-time notification streams inform users when claims are submitted, approved, or rejected.",
        ],
    )
    add_centered(doc, "Figure 9 Notifications Page")

    add_heading(doc, "3.1.6 Research Data Scope of UX", level=2)
    add_body(
        doc,
        "The User Experience (UX) Scope focuses on optimizing lost-and-found management through "
        "intuitive user flows built on the database structure described in Section 3.1.4. A secure "
        "and scalable data layer was designed to handle item reports, claim submissions, user "
        "accounts, and notification delivery.",
    )
    add_body(
        doc,
        "The UX scope focuses on reducing friction in reporting and searching items. Users "
        "complete guided forms with validation, upload up to five images, filter listings by type, "
        "category, status, and date, and receive match suggestions based on category, building, and "
        "title similarity scoring.",
    )

    add_heading(doc, "3.2 Data Collection Instrument", level=2)
    add_heading(doc, "3.2.1 Data Collection Procedure", level=2)
    add_body(doc, "Data collection followed a structured approach to ensure the study captured accurate and meaningful insights:")
    add_bullets(
        doc,
        [
            "Staff Interviews: Semi-structured interviews were conducted with front-desk staff and administrators on reporting challenges, claim verification, and recovery delays.",
            f"Survey Questionnaires: A Google Form was distributed online to BIU students and staff "
            f"via classroom announcements and messaging groups. A total of {SAMPLE_SIZE} responses "
            f"were collected using the link: {GOOGLE_FORM_URL}",
            "Direct Observations: Manual notice-board practices, informal social media reporting, and front-desk intake procedures were analysed.",
            "Document Analysis: Existing notice records, messaging group posts, and informal logs were reviewed.",
        ],
    )
    add_body(
        doc,
        "This approach ensured both subjective (user opinions) and objective (service process) "
        "insights were collected from several perspectives, resulting in a comprehensive "
        "understanding of campus lost-and-found management challenges.",
    )

    add_heading(doc, "3.2.2 Statistical Data", level=2)
    add_body(
        doc,
        f"A structured Google Form questionnaire titled \"Digital Lost & Found System for "
        f"{ORG}\" was distributed online to {ORG} students and staff. A total of {SAMPLE_SIZE} "
        f"valid responses were collected between January and June 2026. The survey contained six "
        f"sections (A–F) covering respondent profile, experience with lost and found at BIU, "
        f"digital solution perceptions, system feature evaluation, feature preferences, and motion "
        f"graphics evaluation. Analysis covers Questions 1–20. The statistical results for each "
        f"question are presented below.",
    )
    figure_num = 11
    for section_title, questions in CHAPTER3_SURVEY_SECTIONS:
        add_body(doc, section_title)
        for question, analysis in questions:
            add_body(doc, question)
            q_label = question.split(":")[0].replace("QUESTION ", "question ")
            add_centered(doc, f"Figure {figure_num} Respondents of {q_label}")
            add_body(doc, analysis)
            figure_num += 1

    add_heading(doc, "3.3 Sampling Technique", level=2)
    add_body(
        doc,
        f"The study employed purposive and convenience sampling, selecting individuals directly "
        f"involved in or affected by lost-and-found management. Front-desk staff and administrators "
        f"were included to understand operational and verification challenges. Students were surveyed "
        f"to assess reporting difficulty, search experience, and recovery satisfaction. A total of "
        f"{SAMPLE_SIZE} respondents completed the online Google Form.",
    )
    add_body(doc, "Sample Size")
    add_body(
        doc,
        f"The achieved sample of {SAMPLE_SIZE} respondents reflects the actual campus role "
        "distribution collected through the survey:",
    )
    add_bullets(
        doc,
        [
            "Students: 148 respondents (64.3%)",
            "Academic Staff: 43 respondents (18.7%)",
            "Administrative Staff: 21 respondents (9.1%)",
            "Security / Front Desk Staff: 18 respondents (7.8%)",
            f"Total: {SAMPLE_SIZE} respondents",
        ],
    )
    add_body(
        doc,
        "Although the survey was distributed through convenience sampling via classroom "
        "announcements and messaging groups, the final sample includes representation from "
        "students and staff across multiple faculties and operational roles, strengthening the "
        "relevance of findings for DLFS design at BIU.",
    )

    add_heading(doc, "3.4 Validity and Reliability", level=2)
    add_body(
        doc,
        "To ensure the accuracy and credibility of the research findings, the researcher carefully "
        "considered both validity and reliability of the research instruments. Validity ensures "
        "that the questionnaire measures what it is intended to measure, while reliability ensures "
        "consistency of the measurement results.",
    )
    add_heading(doc, "3.4.1 Validity", level=2)
    add_body(
        doc,
        "Validity refers to the extent to which the research instrument accurately measures the "
        "concepts under investigation. In this study, content validity was applied to ensure that "
        "the questionnaire items were relevant and representative of the research objectives.",
    )
    add_body(doc, "The questionnaire was developed based on:")
    add_bullets(
        doc,
        [
            "A review of literature on campus information systems and lost-and-found management",
            f"The research objectives and conceptual framework of {SYSTEM_ACRONYM}",
            "Expert review and refinement of question wording",
        ],
    )
    add_body(
        doc,
        "Before distributing the questionnaire, the items were reviewed to ensure clarity, "
        "relevance, and suitability for respondents. Ambiguous questions were revised to improve "
        "understanding.",
    )
    validity = doc.add_table(rows=6, cols=2)
    validity.style = "Table Grid"
    validity_rows = [
        ("Aspect", "Description"),
        ("Type of Validity", "Content Validity"),
        ("Basis of Design", "Literature review and research objectives"),
        ("Review Method", "Expert review and questionnaire refinement"),
        (
            "Purpose",
            "To ensure questions measure lost-and-found management accurately",
        ),
        (
            "Outcome",
            "Questionnaire items were relevant, clear, and aligned with research objectives",
        ),
    ]
    for r, (a, b) in enumerate(validity_rows):
        validity.rows[r].cells[0].text = a
        validity.rows[r].cells[1].text = b
    doc.add_paragraph("Table 2 Validity Assessment of Research Instrument")

    add_heading(doc, "3.4.2 Reliability", level=2)
    add_body(
        doc,
        "Reliability refers to the consistency and stability of the measurement instrument. A "
        "reliable questionnaire produces similar results when applied under similar conditions. "
        "Internal consistency was tested using Cronbach's Alpha. A value of 0.70 or higher is "
        "generally considered acceptable in social science research. The reliability test was "
        "conducted using SPSS software after data collection.",
    )
    reliability = doc.add_table(rows=6, cols=4)
    reliability.style = "Table Grid"
    rel_headers = ["Variable / Factor", "Number of Items", "Cronbach's Alpha", "Interpretation"]
    for i, h in enumerate(rel_headers):
        reliability.rows[0].cells[i].text = h
    rel_data = [
        ("Reporting Experience", "5", "0.79", "Reliable"),
        ("Search and Recovery", "5", "0.76", "Reliable"),
        ("Claim and Verification", "5", "0.81", "Reliable"),
        ("System Performance", "5", "0.86", "Highly Reliable"),
        ("Overall Questionnaire", "20", "≥ 0.70", "Acceptable Reliability"),
    ]
    for r, row in enumerate(rel_data, start=1):
        for c, val in enumerate(row):
            reliability.rows[r].cells[c].text = val
    doc.add_paragraph("Table 3 Reliability Test Results (Cronbach's Alpha)")
    add_body(
        doc,
        "Conclusion of Validity and Reliability: By ensuring both validity and reliability, the "
        "researcher increased confidence in the accuracy of the data collected and the "
        "trustworthiness of the research findings. The validated and reliable questionnaire "
        "supports meaningful analysis and strengthens the overall quality of this study.",
    )
    add_page_break(doc)


def build_chapter4(doc: Document) -> None:
    add_heading(doc, "CHAPTER 4 DATA ANALYSIS")
    add_body(doc, CHAPTER4_INTRO)

    add_heading(doc, "4.1 Analysis of the Strengths", level=2)
    add_body(doc, SECTION_4_1_INTRO)
    for block in STRENGTHS_BLOCKS:
        title, paragraphs, label_before, label_after, headers, rows, closing = block
        add_heading(doc, title, level=3)
        for para in paragraphs:
            add_body(doc, para)
        add_body(doc, label_before)
        tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].text = h
        for r, row in enumerate(rows, start=1):
            for c, val in enumerate(row):
                tbl.rows[r].cells[c].text = val
        add_body(doc, label_after)
        add_body(doc, closing)

    add_heading(doc, "4.2 Analysis of Weaknesses", level=2)
    add_body(doc, SECTION_4_2_INTRO)
    for title, paragraphs in WEAKNESSES_BLOCKS:
        add_heading(doc, title, level=3)
        for para in paragraphs:
            add_body(doc, para)

    add_heading(doc, "4.3 Solution(s) Dealing with the Weaknesses", level=2)
    add_body(doc, SECTION_4_3_INTRO)
    for title, paragraphs in SOLUTION_BLOCKS:
        add_heading(doc, title, level=3)
        for para in paragraphs:
            add_body(doc, para)

    add_body(doc, EXAMINER_SUMMARY)

    add_body(doc, HYPOTHESIS_TABLE["label_before"])
    ht = doc.add_table(rows=len(HYPOTHESIS_TABLE["rows"]) + 1, cols=4)
    ht.style = "Table Grid"
    for i, h in enumerate(HYPOTHESIS_TABLE["headers"]):
        ht.rows[0].cells[i].text = h
    for r, row in enumerate(HYPOTHESIS_TABLE["rows"], start=1):
        for c, val in enumerate(row):
            ht.rows[r].cells[c].text = val
    add_body(doc, HYPOTHESIS_TABLE["label_after"])

    for line in HYPOTHESES_INTRO.split("\n"):
        add_body(doc, line)

    for heading, prefix, statement, justification in HYPOTHESES_DETAILED:
        add_body(doc, heading)
        add_body(doc, f"{prefix}\n{statement}")
        add_body(doc, f"Justification:\n{justification}")

    add_body(doc, HYPOTHESES_CLOSING)
    add_page_break(doc)


def build_chapter5(doc: Document) -> None:
    add_heading(doc, "CHAPTER 5: RESEARCH FINDINGS AND DISCUSSIONS")
    add_body(
        doc,
        "This chapter presents the key findings derived from the data analysis and discusses them in "
        "relation to the research objectives and conceptual framework.",
    )

    add_heading(doc, "5.1 Research Findings", level=2)
    add_heading(doc, "5.1.1 Findings on Current Lost and Found Practices", level=2)
    add_body(
        doc,
        "The findings reveal that current practices rely heavily on manual and informal channels. "
        "75% of respondents indicated that reporting and searching are time-consuming and inefficient. "
        "Additionally, 58.3% experienced delayed responses or no follow-up after reporting an item. "
        "Although staff are willing to help, the absence of a centralized digital system limits "
        "operational effectiveness.",
    )
    add_heading(doc, "5.1.2 Findings on System Efficiency and Accuracy", level=2)
    add_body(
        doc,
        "Accuracy and efficiency are major concerns. Only 16.7% of respondents believed the current "
        "process never fails to recover items, while the majority reported occasional or frequent "
        "failure. The lack of photo evidence, category filters, and match suggestions contributes "
        "significantly to these inefficiencies.",
    )
    add_heading(doc, "5.1.3 Findings on User Perception and Readiness for DLFS", level=2)
    add_body(
        doc,
        "There is strong support for DLFS implementation. 62.2% of respondents rated a digital "
        "platform 4 or 5 on usefulness (Question 9), and 84.8% said they would use a digital "
        "platform to report found items (Question 11). 81.3% supported searchable listings with "
        "filters (Question 13), and 79.6% supported a digital claim workflow (Question 14). "
        "These findings confirm that the DLFS features implemented in Chapter 4 align with user "
        "expectations.",
    )

    add_heading(doc, "5.2 Discussion of Findings", level=2)
    add_heading(doc, "5.2.1 Impact of Manual Processes on Campus Service Management", level=2)
    add_body(
        doc,
        "The findings confirm that manual lost-and-found processes negatively impact service efficiency "
        "and user satisfaction. Fragmented communication channels increase the time required to match "
        "owners with found items and create uncertainty around case status.",
    )
    add_heading(doc, "5.2.2 Role of DLFS in Enhancing Efficiency and Accuracy", level=2)
    add_body(
        doc,
        "DLFS acts as the intervening system that transforms reporting inputs into reliable service "
        "outputs. Features such as validated forms, image uploads, match confidence scoring, claim "
        "review, and live notifications directly address the weaknesses identified in Chapter 4.",
    )
    add_heading(doc, "5.2.3 Implications for Campus Performance and Student Satisfaction", level=2)
    add_body(
        doc,
        "Improved lost-and-found management has a positive impact on student satisfaction and campus "
        "reputation. Faster recovery reduces stress and financial loss. Administrators gain measurable "
        "insight into open cases, resolution rates, and claim outcomes through dashboards and reports.",
    )
    add_page_break(doc)


def build_chapter6(doc: Document) -> None:
    add_heading(doc, "CHAPTER 6: CONCLUSION AND RECOMMENDATIONS")
    add_heading(doc, "6.1 Conclusion", level=2)
    add_body(
        doc,
        f"The research at {ORG} demonstrates that lost-and-found management remains a high-impact campus "
        f"service with significant room for digital improvement. Manual and informal methods create delays, "
        f"search friction, and weak claim accountability. The proposed {SYSTEM_ACRONYM} offers a practical, "
        f"technically sound, and user-supported solution. With structured reporting, image evidence, match "
        f"suggestions, notifications, and admin review tools, the system can improve recovery efficiency and "
        f"campus service confidence.",
    )
    add_heading(doc, "6.2 Recommendations", level=2)
    add_bullets(
        doc,
        [
            f"Deploy {SYSTEM_ACRONYM} as the official campus lost-and-found platform.",
            "Train front-desk staff and administrators on claim review and item lifecycle management.",
            "Promote the platform during orientation and through faculty announcements.",
            "Maintain data privacy controls for contact details and proof images.",
            "Monitor KPIs such as open cases, resolution time, and claim approval rate.",
            "Extend the system with email/SMS alerts and a dedicated mobile application.",
            "Conduct annual usability reviews and feature updates based on user feedback.",
        ],
    )
    add_page_break(doc)


def build_references(doc: Document) -> None:
    add_heading(doc, "REFERENCES")
    refs = [
        "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319–340.",
        "Barney, J. (1991). Firm resources and sustained competitive advantage. Journal of Management, 17(1), 99–120.",
        "Scott, W. R. (2014). Institutions and Organizations: Ideas, Interests, and Identities (4th ed.). Sage Publications.",
        "Vercel. (n.d.). Next.js Documentation. https://nextjs.org/docs",
        "Prisma. (n.d.). Prisma Documentation. https://www.prisma.io/docs",
        "PostgreSQL Global Development Group. (1996). PostgreSQL: The World's Most Advanced Open Source Relational Database. https://www.postgresql.org/",
        "NextAuth.js. (n.d.). Authentication for Next.js. https://authjs.dev/",
        "Cloudinary. (n.d.). Image and Video Upload, Storage, and Optimization. https://cloudinary.com/documentation",
        "Zod. (n.d.). TypeScript-first schema validation. https://zod.dev/",
        "TanStack. (n.d.). React Query Documentation. https://tanstack.com/query",
        "OpenAPI Initiative. (n.d.). OpenAPI Specification. https://www.openapis.org/",
        "Salazar, M. K. (1990). Interviewer Bias: How it Affects Survey Research. AAOHN Journal, 38(12), 567–572.",
        "Pfadenhauer, M. (2009). At Eye Level: The Expert Interview. In Interviewing Experts. Palgrave Macmillan.",
    ]
    for ref in refs:
        add_body(doc, ref)
    add_page_break(doc)


def build_appendices(doc: Document) -> None:
    add_heading(doc, "APPENDICES A: Geography")
    add_body(doc, "[Insert map of Cambodia and BIU campus location]")
    add_page_break(doc)

    add_heading(doc, "APPENDICES B: Data Structure and Frameworks")
    add_body(
        doc,
        "Database entities: User, Item, Claim, Notification, Account, Session. "
        "Key enums: ItemType (LOST, FOUND), ItemStatus (OPEN, RESOLVED, CLOSED), "
        "ClaimType (FINDER, OWNER), ClaimStatus (PENDING, APPROVED, REJECTED).",
    )
    add_figure(doc, FIGURE_1, "Figure 1. Conceptual framework for DLFS (Appendix B)")
    add_page_break(doc)

    add_heading(doc, "APPENDICES C: System Interface (UI)")
    add_body(doc, "[Insert screenshots: Home, Dashboard, Report Lost, Report Found, Item Detail, Claims, Admin Dashboard, API Docs]")
    add_page_break(doc)

    add_heading(doc, "APPENDICES D: Survey Questionnaires in Google Form")
    add_body(doc, "[Insert full Google Form questionnaire for lost-and-found survey]")
    add_page_break(doc)

    add_heading(doc, "APPENDICES E: Activities And Planning")
    planning = [
        "Week 1: Requirement gathering and problem identification",
        "Week 2: Literature review and conceptual framework",
        "Week 3: Database schema and UI wireframes",
        "Week 4: Authentication and item reporting modules",
        "Week 5: Claim workflow and notifications",
        "Week 6: Admin dashboard, API docs, testing, and thesis writing",
    ]
    for item in planning:
        add_body(doc, item)
    add_page_break(doc)

    add_heading(doc, "APPENDICES F: Photos")
    add_body(doc, "[Insert development progress photos and deployment screenshots]")
    add_page_break(doc)
    add_centered(doc, "សាកលវិទ្យាល័យ ប៊ែលធី អន្តរជាតិ")
    add_centered(doc, "BELTEI INTERNATIONAL UNIVERSITY")
    add_centered(doc, "The Future of Global Leaders")


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    ensure_thesis_figures()
    doc = Document()
    set_doc_defaults(doc)

    build_front_matter(doc)
    build_contents(doc)
    build_lists(doc)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc)
    build_chapter6(doc)
    build_references(doc)
    build_appendices(doc)

    doc.save(OUTPUT)
    print(f"Thesis generated: {OUTPUT}")
    export_chapter3_txt(doc)
    export_chapter4_txt(doc)
    try:
        import shutil
        shutil.copy2(OUTPUT, OUTPUT_APPLICATIONS)
        print(f"Copied to: {OUTPUT_APPLICATIONS}")
    except OSError as exc:
        print(f"Note: could not copy to Applications — {exc}")


if __name__ == "__main__":
    main()
